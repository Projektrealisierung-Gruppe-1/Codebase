from fpdf import FPDF
from docx import Document

def pdfcreator(txt,zsm_txt,kltxt,senttxt):
    # variable pdf
    pdf = FPDF()
    
    # Add a page
    pdf.add_page()
    
    # set style and size of font
    # that you want in the pdf
    pdf.set_font("Arial", "B", size = 18)
    # create a cell
    pdf.cell(200, 10, txt = "Projektrealisierung Download-Bericht",
            ln = 1, align = 'C')

    
    pdf.set_font("Arial", "U", size = 14)
    pdf.cell(200, 10, txt = "Text Zusammenfassung",
            ln = 1, align = 'L')
    pdf.set_font("Arial", size = 12)
    # add another cell
    pdf.multi_cell(200, 5, txt = zsm_txt, align = 'L')

    pdf.set_font("Arial","U", size = 14)
    pdf.cell(200, 15, txt = "Text Klassifizierung",
            ln = 1, align = 'L')
    pdf.set_font("Arial", size = 12)
    pdf.multi_cell(200, 5, txt = kltxt, align = 'L')


    pdf.set_font("Arial","U", size = 14)
    pdf.cell(200, 15, txt = "Text Sentiment",
            ln = 1, align = 'L')
    pdf.set_font("Arial", size = 12)
    pdf.multi_cell(200, 5, txt = senttxt, align = 'L')
    
    pdf.add_page()
    pdf.set_font("Arial", "U",size = 14)
    pdf.cell(200, 10, txt = "Original Text",
            ln = 1, align = 'L')
    pdf.set_font("Arial", size = 12)
    # add another cell
    pdf.multi_cell(200, 5, txt = txt, align = 'L')
    
    
    pdf.output("download.pdf")
    # save the pdf with name .pdf

def docxcreator(txt,zsm_txt,kltxt,senttxt):

        document = Document()
        document.add_heading('Projektrealisierung Download-Bericht', 0)


        document.add_heading('Text Zusammenfassung', level=1)
        document.add_paragraph(zsm_txt)
        document.add_heading('Text Klassifizierung', level=1)
        document.add_paragraph(kltxt)
        document.add_heading('Text Sentiment', level=1)
        document.add_paragraph(senttxt)

        document.add_page_break()
        document.add_heading('Original Text', level=1)
        document.add_paragraph(txt)
        document.save('download.docx')

        return document

